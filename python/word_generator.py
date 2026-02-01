#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Word文档生成模块
"""

import sys
import json
import os
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from PIL import Image


def split_chinese_english(text):
    """
    将文本分割为中文和英文部分 - ruanzhu19项目的原始实现
    返回列表，每个元素为(文本片段, 是否为中文)的元组
    """
    if not text:
        return []

    result = []
    current_text = ""
    current_is_chinese = None

    for char in text:
        is_chinese = '\u4e00' <= char <= '\u9fff'

        if current_is_chinese is not None and current_is_chinese != is_chinese:
            if current_text:
                result.append((current_text, current_is_chinese))
            current_text = char
            current_is_chinese = is_chinese
        else:
            current_text += char
            if current_is_chinese is None:
                current_is_chinese = is_chinese

    if current_text:
        result.append((current_text, current_is_chinese))

    return result


def clean_markdown_content(content):
    """
    清理Markdown格式，保留基本结构
    """
    if not content:
        return ""
    
    content = re.sub(r'^软件名称：.*?\n', '', content, flags=re.MULTILINE)
    content = re.sub(r'^版本号：.*?\n', '', content, flags=re.MULTILINE)
    content = re.sub(r'^开发完成日期：.*?\n', '', content, flags=re.MULTILINE)
    
    content = re.sub(r'```[\s\S]*?```', '', content)
    content = re.sub(r'`([^`]+)`', r'\1', content)
    
    lines = content.split('\n')
    processed_lines = []
    seen_titles = set()
    
    for line in lines:
        line_stripped = line.strip()
        
        if re.match(r'^#+\s*(.+)$', line_stripped):
            title = re.sub(r'^#+\s*', '', line_stripped).strip()
            if title in seen_titles:
                continue
            seen_titles.add(title)
            processed_lines.append(title)
        elif re.match(r'^第[一二三四五六七八九十\d]+章', line_stripped):
            if line_stripped in seen_titles:
                continue
            seen_titles.add(line_stripped)
            processed_lines.append(line)
        else:
            processed_lines.append(line)
    
    content = '\n'.join(processed_lines)
    
    content = re.sub(r'\*\*(.+?)\*\*', r'\1', content)
    content = re.sub(r'\*(.+?)\*', r'\1', content)
    
    content = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', content)
    
    content = re.sub(r'^[\*\-\+]\s+', '• ', content, flags=re.MULTILINE)
    content = re.sub(r'^\d+\.\s+', '', content, flags=re.MULTILINE)
    
    content = re.sub(r'\n{3,}', '\n\n', content)
    
    return content.strip()


def setup_header_footer(doc, app_name, doc_type="manual"):
    """
    设置页眉页脚 - 完全基于ruanzhu19项目的原始实现
    doc_type: "manual"=说明书, "source"=源代码, "info"=软著信息
    """
    section = doc.sections[0]
    
    header = section.header
    header_para = header.paragraphs[0]

    if doc_type == "source":
        setup_ruanzhu19_header(header_para, app_name)
    elif doc_type == "info":
        header_para.text = f"{app_name} 软件著作权申请信息"
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        setup_ruanzhu19_header(header_para, app_name)

    footer = section.footer
    footer_para = footer.paragraphs[0]

    if doc_type == "manual":
        setup_ruanzhu19_footer(footer_para)
    else:
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER


def setup_ruanzhu19_header(header_para, app_name):
    """
    设置ruanzhu19项目的原始页眉格式：
    - 应用名称 V1.0 (居中)
    - 第 X 页 (右对齐)
    - 页眉下划线
    """
    tab_stops = header_para.paragraph_format.tab_stops
    for stop in tab_stops:
        tab_stops.remove(stop)
    tab_stops.add_tab_stop(Inches(3.25), WD_TAB_ALIGNMENT.CENTER)
    tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)

    run = header_para.add_run('\t')

    title_text = f"{app_name} V1.0"
    run = header_para.add_run(title_text)
    run.font.size = Pt(12)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    run = header_para.add_run('\t')

    run = header_para.add_run('第 ')
    run.font.size = Pt(12)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.text = ' PAGE '
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    
    run = header_para.add_run(' 页')
    run.font.size = Pt(12)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    pPr = header_para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)


def setup_ruanzhu19_footer(footer_para):
    """
    设置ruanzhu19项目的原始页脚格式：
        - 居中页码
    """
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.text = ' PAGE '
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')

    run = footer_para.add_run()
    run.font.size = Pt(12)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


def format_document_with_proper_headings(content, app_name, clean_markdown=False):
    """
    格式化文档内容 - 基于ruanzhu19项目，不添加Markdown标记
    clean_markdown: 是否清理Markdown格式（仅用于软著申请信息）
    """
    if not content:
        return ""

    if clean_markdown:
        content = clean_markdown_content(content)

    return content


def prepare_screenshots_for_sections(screenshots_dir="screenshots", app_name=None):
    """
    准备截图列表，返回5张截图的路径，用于插入到第三章的5个小节中
    返回格式: [screenshot1_path, screenshot2_path, ...] 或空列表
    """
    if not screenshots_dir or not os.path.exists(screenshots_dir):
        return []

    all_screenshot_files = []
    for file in os.listdir(screenshots_dir):
        if file.lower().endswith('.png'):
            all_screenshot_files.append(file)

    if not all_screenshot_files:
        return []

    screenshot_files = []
    if app_name:
        clean_app_name = re.sub(r'[^a-zA-Z0-9\u4e00-\u9fa5]', '_', app_name)
        for file in all_screenshot_files:
            if file.startswith(clean_app_name):
                screenshot_files.append(file)
        
        if not screenshot_files:
            file_times = []
            for file in all_screenshot_files:
                file_path = os.path.join(screenshots_dir, file)
                mtime = os.path.getmtime(file_path)
                file_times.append((mtime, file))
            file_times.sort(reverse=True)
            screenshot_files = [file for _, file in file_times[:5]]
    else:
        screenshot_files = all_screenshot_files

    if not screenshot_files:
        return []

    screenshot_files = sorted(screenshot_files)[:5]
    
    return [os.path.join(screenshots_dir, file) for file in screenshot_files]


def insert_single_screenshot(doc, screenshot_path, section_title, figure_num=None):
    """
    在文档中插入单张截图
    section_title: 截图所属的小节标题，如 "首页"
    figure_num: 图序号，如 1、2、3...（可选）
    如果截图不存在或加载失败，直接跳过，不插入任何内容
    """
    if not screenshot_path or not os.path.exists(screenshot_path):
        return

    try:
        image_para = doc.add_paragraph()
        image_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        with Image.open(screenshot_path) as img:
            width, height = img.size
            max_width = Inches(6)
            if width > height:
                new_width = max_width
                new_height = Inches(height * 6 / width)
            else:
                new_height = Inches(6)
                new_width = Inches(width * 6 / height)

        run = image_para.add_run()
        with open(screenshot_path, 'rb') as f:
            run.add_picture(f, width=new_width, height=new_height)

        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        if figure_num is not None:
            title_text = f'图{figure_num}：{section_title}界面'
        else:
            title_text = f'图：{section_title}界面'
        
        title_run = title_para.add_run(title_text)
        title_run.font.size = Pt(12)
        title_run.font.name = '宋体'
        title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        title_run.bold = True

        doc.add_paragraph('')
        
    except Exception as e:
        pass


def insert_screenshot_images(doc, content, screenshots_dir="screenshots", app_name=None):
    """
    在文档中插入截图图片（已废弃，保留以防兼容性问题）
    """
    if not os.path.exists(screenshots_dir):
        # sys.stderr.flush()
        return

    all_screenshot_files = []
    for file in os.listdir(screenshots_dir):
        if file.lower().endswith('.png'):
            all_screenshot_files.append(file)

    if not all_screenshot_files:
        # sys.stderr.flush()
        return

    screenshot_files = []
    if app_name:
        clean_app_name = re.sub(r'[^a-zA-Z0-9\u4e00-\u9fa5]', '_', app_name)
        # sys.stderr.flush()
        for file in all_screenshot_files:
            if file.startswith(clean_app_name):
                screenshot_files.append(file)
            # else:
        # sys.stderr.flush()
        
        if not screenshot_files:
            file_times = []
            for file in all_screenshot_files:
                file_path = os.path.join(screenshots_dir, file)
                mtime = os.path.getmtime(file_path)
                file_times.append((mtime, file))
            
            file_times.sort(reverse=True)
            screenshot_files = [file for _, file in file_times[:5]]
    else:
        screenshot_files = all_screenshot_files

    if not screenshot_files:
        # sys.stderr.flush()
        return

    doc.add_page_break()

    appendix_para = doc.add_paragraph()
    appendix_run = appendix_para.add_run('附录：系统界面截图')
    appendix_run.font.name = '黑体'
    appendix_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    appendix_run.font.size = Pt(16)
    appendix_run.bold = True
    appendix_run.font.color.rgb = RGBColor(0, 0, 0)
    appendix_para.paragraph_format.space_before = Pt(13)
    appendix_para.paragraph_format.space_after = Pt(13)

    desc_para = doc.add_paragraph()
    desc_para.paragraph_format.first_line_indent = Pt(24)
    desc_run = desc_para.add_run(f'本系统的主要界面截图如下所示，共包含 {len(screenshot_files)} 张截图，展示了系统的核心功能和用户交互界面。')
    desc_run.font.size = Pt(12)
    desc_run.font.name = '宋体'
    desc_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    for i, screenshot_file in enumerate(sorted(screenshot_files), 1):
        screenshot_path = os.path.join(screenshots_dir, screenshot_file)

        try:
            title_para = doc.add_paragraph()
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = title_para.add_run(f'图 {i}：系统界面截图')
            title_run.font.size = Pt(12)
            title_run.font.name = '宋体'
            title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            title_run.bold = True

            image_para = doc.add_paragraph()
            image_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            with Image.open(screenshot_path) as img:
                width, height = img.size
                max_width = Inches(6)
                if width > height:
                    new_width = max_width
                    new_height = Inches(height * 6 / width)
                else:
                    new_height = Inches(6)
                    new_width = Inches(width * 6 / height)

            run = image_para.add_run()
            with open(screenshot_path, 'rb') as f:
                run.add_picture(f, width=new_width, height=new_height)

            doc.add_paragraph('')

            
        except Exception as e:
            error_para = doc.add_paragraph()
            error_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            error_run = error_para.add_run(f'图 {i}：截图加载失败 ({screenshot_file})')
            error_run.font.size = Pt(12)
            error_run.font.name = '宋体'
            error_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            error_run.font.color.rgb = RGBColor(255, 0, 0)

            doc.add_paragraph('')


def create_software_info_document(app_name, content, output_path):
    """
    创建软著申请信息Word文档 - 基于ruanzhu19项目的简化实现
    """
    doc = Document()

    section = doc.sections[0]
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    info = parse_software_info_content(content, app_name)

    title = doc.add_heading(f"{info['clean_name']} 软件著作权申请信息", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading('基本信息', level=1)

    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'

    header_cells = table.rows[0].cells
    header_cells[0].text = '项目'
    header_cells[1].text = '内容'

    row_cells = table.add_row().cells
    row_cells[0].text = '软件名称'
    row_cells[1].text = info['clean_name']

    row_cells = table.add_row().cells
    row_cells[0].text = '版本号'
    row_cells[1].text = info['version']

    row_cells = table.add_row().cells
    row_cells[0].text = '开发的硬件环境'
    row_cells[1].text = info['dev_hardware']

    row_cells = table.add_row().cells
    row_cells[0].text = '运行的硬件环境'
    row_cells[1].text = info['runtime_hardware']

    row_cells = table.add_row().cells
    row_cells[0].text = '开发该软件的操作系统'
    row_cells[1].text = info['dev_os']

    row_cells = table.add_row().cells
    row_cells[0].text = '软件开发环境/开发工具'
    row_cells[1].text = info['dev_tools']

    row_cells = table.add_row().cells
    row_cells[0].text = '该软件的运行平台/操作系统'
    row_cells[1].text = info['runtime_os']

    row_cells = table.add_row().cells
    row_cells[0].text = '软件运行支撑环境/支持软件'
    row_cells[1].text = info['runtime']

    row_cells = table.add_row().cells
    row_cells[0].text = '编程语言'
    row_cells[1].text = info['languages']

    row_cells = table.add_row().cells
    row_cells[0].text = '源程序量'
    row_cells[1].text = info['code_lines']

    row_cells = table.add_row().cells
    row_cells[0].text = '开发目的'
    row_cells[1].text = info['purpose']

    row_cells = table.add_row().cells
    row_cells[0].text = '面向领域/行业'
    row_cells[1].text = info['domain']

    row_cells = table.add_row().cells
    row_cells[0].text = '软件的主要功能'
    row_cells[1].text = info['functions']

    row_cells = table.add_row().cells
    row_cells[0].text = '软件的技术特点'
    row_cells[1].text = info['features']

    doc.save(output_path)
    return output_path


def parse_software_info_content(content, app_name):
    """
    解析Java端生成的软著信息结构化文本内容
    """
    clean_name = app_name
    if '-软著申请表' in clean_name:
        clean_name = clean_name.replace('-软著申请表', '')
    if '软著申请表' in clean_name:
        clean_name = clean_name.replace('软著申请表', '')
    
    info = {
        'name': app_name,
        'clean_name': clean_name,
        'version': '',
        'dev_hardware': '',
        'runtime_hardware': '',
        'dev_os': '',
        'runtime_os': '',
        'dev_tools': '',
        'runtime': '',
        'languages': '',
        'code_lines': '',
        'purpose': '',
        'domain': '',
        'functions': '',
        'features': '',
        'database': ''
    }

    if not content:
        return info


    content = content.replace('\r\n', '\n').replace('\r', '\n')

    name_match = re.search(r'软件名称：(.+)', content)
    if name_match:
        info['name'] = name_match.group(1).strip()

    version_match = re.search(r'版本号：(.+)', content)
    if version_match:
        info['version'] = version_match.group(1).strip()

    purpose_found = False

    purpose_match = re.search(r'开发目的：(.+?)(?=\n[一二三四五六七八九十]|技术特点：|主要功能：|面向领域|$)', content, re.DOTALL)
    if purpose_match:
        purpose_text = purpose_match.group(1).strip()
        if purpose_text and len(purpose_text) > 5:
            info['purpose'] = purpose_text
            purpose_found = True
    
    if not purpose_found:
        purpose_match = re.search(r'软件用途：(.+?)(?=\n[一二三四五六七八九十]|技术特点：|主要功能：|面向领域|$)', content, re.DOTALL)
        if purpose_match:
            purpose_text = purpose_match.group(1).strip()
            if purpose_text and len(purpose_text) > 5:
                info['purpose'] = purpose_text
                purpose_found = True

    if not purpose_found:
        section_match = re.search(r'二、软件用途和技术特点\s*\n\n(.+?)(?=\n\n[一二三四五六七八九十]|$)', content, re.DOTALL)
        if section_match:
            section_content = section_match.group(1).strip()
            if '软件用途：' in section_content:
                lines = section_content.split('\n')
                purpose_lines = []
                in_purpose_section = False
                for line in lines:
                    if '软件用途：' in line:
                        in_purpose_section = True
                        after_colon = line.split('软件用途：')[1].strip()
                        if after_colon:
                            purpose_lines.append(after_colon)
                    elif in_purpose_section:
                        if '技术特点：' in line or '主要功能：' in line:
                            break
                        if line.strip():
                            purpose_lines.append(line.strip())

                if purpose_lines:
                    purpose_text = '\n'.join(purpose_lines).strip()
                    if len(purpose_text) > 10:
                        info['purpose'] = purpose_text
                        purpose_found = True

    if not purpose_found:
        section_match = re.search(r'二、软件用途和技术特点\s*\n\n(.+?)(?=技术特点：|主要功能：|\n\n[一二三四五六七八九十]|$)', content, re.DOTALL)
        if section_match:
            first_part = section_match.group(1).strip()
            first_part = re.sub(r'^软件用途：\s*', '', first_part)
            if first_part and len(first_part) > 10 and '技术特点' not in first_part:
                info['purpose'] = first_part
                purpose_found = True

    features_match = re.search(r'技术特点：([^。]+?)(?=。|编程语言|开发工具|运行环境|$)', content)
    if not features_match:
        features_match = re.search(r'技术特点：(.+?)(?=编程语言|开发工具|运行环境|$)', content, re.DOTALL)
    if features_match:
        features_text = features_match.group(1).strip()
        features_text = re.sub(r'\s+', '，', features_text)
        info['features'] = features_text

    functions_found = False

    functions_match = re.search(r'主要功能：([^。]+?)(?=。|技术特点|编程语言|开发工具|$)', content)
    if functions_match:
        functions_text = functions_match.group(1).strip()
        functions_text = functions_text.replace('、', '，')
        if functions_text and len(functions_text) > 5:
            info['functions'] = functions_text
            functions_found = True

    if not functions_found:
        functions_match = re.search(r'主要功能：(.+?)(?=\n[一二三四五六七八九十]|技术特点|编程语言|开发工具|$)', content, re.DOTALL)
        if functions_match:
            functions_text = functions_match.group(1).strip()
            functions_text = re.sub(r'\s+', '，', functions_text)
            functions_text = functions_text.replace('、', '，')
            if functions_text and len(functions_text) > 5:
                info['functions'] = functions_text
                functions_found = True

    info['dev_tools'] = "Visual Studio Code"

    info['languages'] = "HTML、JavaScript、Python"

    info['runtime'] = "Web浏览器、Python 3.x"

    info['database'] = "MySQL 8.0"

    domain_match = re.search(r'面向领域/行业：([^。]+?)(?=。|主要功能|技术特点|编程语言|$)', content)
    if domain_match:
        info['domain'] = domain_match.group(1).strip()
    elif not info['domain']:
        info['domain'] = "信息技术服务业"

    code_lines_match = re.search(r'源程序量：(.+)', content)
    if code_lines_match:
        info['code_lines'] = code_lines_match.group(1).strip()
    else:
        lines_patterns = [
            r'源代码总行数[:\s]*(\d+)',
            r'过滤后的源代码总行数[:\s]*(\d+)', 
            r'源代码行数[:\s]*(\d+)',
            r'代码总行数[:\s]*(\d+)',
            r'总行数[:\s]*(\d+)',
            r'行数[:\s]*(\d+)'
        ]
        
        lines_count = None
        for pattern in lines_patterns:
            lines_match = re.search(pattern, content)
            if lines_match:
                lines_count = lines_match.group(1)
                break
        
        if lines_count:
            info['code_lines'] = f"{lines_count}行"
        else:
            info['code_lines'] = calculate_source_code_lines()

    info['dev_hardware'] = "i5-12400 CPU，16G 内存，1TB 固态硬盘，20Mbps网络带宽。"

    info['runtime_hardware'] = "服务器内存2G及以上，硬盘空间128GB及以上，网络带宽1Mbps及以上。"

    info['dev_os'] = "Windows 11系统。"

    info['runtime_os'] = "Windows 11，Windows 10，Windows Server 2016等桌面操作系统。"

    if info['features']:
        cleaned_features = clean_technical_content(info['features'])
        if cleaned_features:
            info['features'] = cleaned_features
        else:
            info['features'] = "采用现代化系统架构，具有良好的稳定性和扩展性，支持高并发访问，界面友好易用"
    else:
        info['features'] = "采用现代化系统架构，具有良好的稳定性和扩展性，支持高并发访问，界面友好易用"
    
    if info['functions']:
        cleaned_functions = clean_technical_content(info['functions'])
        if cleaned_functions:
            info['functions'] = cleaned_functions

    if not info['domain']:
        domain = infer_domain_from_content(app_name, info['purpose'], info['functions'])
        if domain:
            info['domain'] = domain

    return info


def process_function_list(functions_text):
    """
    处理功能列表，将章节标题转换为功能描述
    """
    if not functions_text:
        return ""

    lines = functions_text.split('\n')
    processed_functions = []
    
    for line in lines:
        line = line.strip()
        if not line:
            continue

        line = re.sub(r'^\d+\.\s*', '', line)

        if re.match(r'
            chapter_match = re.search(r'第[一二三四五六七八九十\d]+章\s*(.+)', line)
            if chapter_match:
                chapter_name = chapter_match.group(1).strip()
                if chapter_name:
                    processed_functions.append(chapter_name)
        elif '第' in line and '章' in line:
            chapter_match = re.search(r'第[一二三四五六七八九十\d]+章\s*(.+)', line)
            if chapter_match:
                chapter_name = chapter_match.group(1).strip()
                if chapter_name:
                    processed_functions.append(chapter_name)
        else:
            if len(line) > 3:
                processed_functions.append(line)

    if processed_functions:
        return '、'.join(processed_functions[:5])
    else:
        return ""


def calculate_source_code_lines(project_root="."):
    """
    计算项目源代码行数（只计算主要业务代码）
    """
    code_extensions = {
        '.java', '.py', '.js', '.ts', '.jsx', '.tsx', '.vue', '.sql'
    }
    
    exclude_dirs = {
        'node_modules', '.git', '__pycache__', 'venv', 'env', 'build', 'dist',
        'target', 'bin', 'obj', '.vs', '.vscode', '.idea', 'logs', 'temp',
        'screenshots', 'images', 'docs', 'documentation', 'test', 'tests',
        'spec', 'mock', 'mocks', 'static', 'assets', 'public', 'resources',
        'lib', 'libs', 'vendor', 'third-party', 'external', 'node-modules',
        'fastgpt', 'onlyoffice', 'opt', 'ruanzhu19', 'temp'
    }
    
    exclude_files = {
        'test', 'spec', 'mock', '.test.', '.spec.', '.mock.', 'node_modules',
        'package-lock.json', 'yarn.lock', 'pom.xml'
    }
    
    total_lines = 0
    file_count = 0
    
    try:
        for root, dirs, files in os.walk(project_root):
            dirs[:] = [d for d in dirs if d not in exclude_dirs and not d.startswith('.')]
            
            if not any(keyword in root.lower() for keyword in ['src', 'source', 'main', 'app', 'lib', 'api', 'biz', 'start']):
                continue
                
            for file in files:
                _, ext = os.path.splitext(file.lower())
                if ext in code_extensions:
                    if any(pattern in file.lower() for pattern in exclude_files):
                        continue
                        
                    file_path = os.path.join(root, file)
                    try:
                        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                            lines = sum(1 for line in f if line.strip() and not line.strip().startswith(('//','
                            if lines > 10:
                                total_lines += lines
                                file_count += 1
                    except:
                        try:
                            with open(file_path, 'r', encoding='gbk', errors='ignore') as f:
                                lines = sum(1 for line in f if line.strip() and not line.strip().startswith(('//','#','*','<!--')))
                                if lines > 10:
                                    total_lines += lines
                                    file_count += 1
                        except:
                            continue
    except Exception as e:
        return "约3000行"
    
    if total_lines > 0:
        return f"约{total_lines}行（包含{file_count}个源文件）"
    else:
        return "约3000行"


def clean_technical_content(content):
    """
    清理技术特点和主要功能内容，移除编程语言和源程序量信息
    """
    if not content:
        return ""
    
    content = re.sub(r'\b(?:Java|Python|JavaScript|HTML|CSS|PHP|C\+\+|C#|Vue\.js|React\.js|Spring Boot|jQuery)\b[，、；：\s]*', '', content)
    
    content = re.sub(r'[，、；：\s]*(?:源程序量|代码行数|程序行数|共\s*\d+\s*行)[^，。]*', '', content)
    
    content = re.sub(r'[，、；：\s]*(?:使用|采用|基于)\s*(?:Visual Studio|IntelliJ|Eclipse|Git|Maven|Gradle)[^，。]*', '', content)
    
    content = re.sub(r'[，、；：\s]+', '，', content)
    content = re.sub(r'^[，、；：\s]+|[，、；：\s]+$', '', content)
    
    return content.strip()


def infer_domain_from_content(app_name, purpose, functions):
    """
    根据应用名称、用途和功能智能推断面向领域
    """
    all_text = f"{app_name} {purpose} {functions}".lower()

    domain_keywords = {
        '电子商务': ['电商', '购物', '商城', '零售', '销售', '营销', '客户', '用户行为', '商品'],
        '金融服务': ['金融', '银行', '支付', '理财', '投资', '保险', '贷款', '风控'],
        '教育培训': ['教育', '学习', '培训', '课程', '学生', '教师', '知识'],
        '医疗健康': ['医疗', '健康', '医院', '患者', '诊断', '治疗', '药品'],
        '物流运输': ['物流', '运输', '配送', '仓储', '供应链', '快递'],
        '制造业': ['生产', '制造', '工厂', '设备', '质量', '工艺'],
        '房地产': ['房地产', '房屋', '物业', '租赁', '买卖'],
        '政务服务': ['政务', '政府', '公共', '行政', '办事'],
        '企业管理': ['企业', '管理', '办公', '人事', '财务', '项目'],
        '数据分析': ['数据', '分析', '统计', '报表', '智能', '算法', '大数据'],
        '社交娱乐': ['社交', '娱乐', '游戏', '聊天', '社区', '内容'],
        '旅游服务': ['旅游', '酒店', '景点', '出行', '预订'],
    }

    domain_scores = {}
    for domain, keywords in domain_keywords.items():
        score = 0
        for keyword in keywords:
            if keyword in all_text:
                score += 1
        domain_scores[domain] = score

    if domain_scores:
        best_domain = max(domain_scores, key=domain_scores.get)
        if domain_scores[best_domain] > 0:
            return best_domain

    return None


def create_manual_document(app_name, content, output_path):
    """
    创建软件说明书Word文档 - 完全基于ruanzhu19项目的原始实现
    """
    formatted_text = format_document_with_proper_headings(content, app_name, clean_markdown=True)

    doc = Document()

    section = doc.sections[0]
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    setup_header_footer(doc, app_name, "manual")

    create_ruanzhu19_cover(doc, app_name)

    screenshot_list = prepare_screenshots_for_sections(screenshot_dir, app_name)

    process_ruanzhu19_content(doc, formatted_text, screenshot_list)

    doc.save(output_path)
    return output_path


def create_ruanzhu19_cover(doc, app_name):
    """
    创建ruanzhu19项目的原始封面格式
    """
    doc.add_paragraph('')
    for _ in range(4):
        doc.add_paragraph('')

    title_para = doc.add_paragraph()
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title_para.add_run('软件说明书')
    title_run.font.size = Pt(48)
    title_run.font.name = '宋体'
    title_run.bold = True
    title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    for _ in range(4):
        doc.add_paragraph('')

    app_para = doc.add_paragraph()
    app_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    app_run = app_para.add_run(f"{app_name} V1.0")
    app_run.font.size = Pt(36)
    app_run.font.name = '宋体'
    app_run.bold = True
    app_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    doc.add_page_break()


def process_ruanzhu19_content(doc, formatted_text, screenshot_list=None):
    """
    处理文档内容 - 使用ruanzhu19的原始格式和样式
    标题下方不空行，行间距1.25，在第三章的5个小节后插入对应截图
    screenshot_list: 截图路径列表，顺序对应 3.2-3.6 小节
    """
    lines = formatted_text.split('\n')

    processed_lines = []
    found_first_appendix = False
    found_second_appendix = False

    for line in lines:
        if line.strip().startswith('附录：系统界面截图'):
            if not found_first_appendix:
                found_first_appendix = True
                continue
            elif not found_second_appendix:
                found_second_appendix = True
                processed_lines.append(line)
            else:
                processed_lines.append(line)
        elif found_first_appendix and not found_second_appendix:
            if line.strip().startswith('本系统的主要界面截图如下所示，展示了系统的核心功能和用户交互界面。'):
                continue
        else:
            processed_lines.append(line)

    section_screenshot_map = {
        '3.2': ('首页', 0),
        '3.3': ('用户管理', 1),
        '3.4': ('数据分析', 2),
        '3.5': ('系统设置', 3),
        '3.6': ('消息中心', 4)
    }
    
    pending_screenshot = None
    
    for i, line in enumerate(processed_lines):
        if not line.strip():
            continue

        line = re.sub(r'(\d+)\s+\.(\d+)', r'\1.\2', line)
        
        section_match = re.match(r'^(3\.[2-6])\s+', line)
        
        is_new_section = (re.match(r'^\d+\.\d+\s+', line) is not None) or (line.strip().startswith('第') and '章' in line)
        
        if pending_screenshot and is_new_section:
            section_title, screenshot_idx = pending_screenshot
            if screenshot_list and screenshot_idx < len(screenshot_list):
                insert_single_screenshot(doc, screenshot_list[screenshot_idx], section_title, screenshot_idx + 1)
            pending_screenshot = None

        paragraph = doc.add_paragraph()

        if line.strip().startswith('第') and '章' in line:
            paragraph.style = 'Heading 1'
            run = paragraph.add_run(line)
            run.font.name = '黑体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            run.font.size = Pt(16)
            run.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)
            paragraph.paragraph_format.space_before = Pt(13)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.25

        elif re.match(r'^\d+\.\d+\.\d+\s+', line):
            paragraph.style = 'Heading 3'
            run = paragraph.add_run(line)
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            run.font.size = Pt(12)
            run.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)
            paragraph.paragraph_format.space_before = Pt(10)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.25

        elif re.match(r'^\d+\.\d+\s+', line):
            paragraph.style = 'Heading 2'
            run = paragraph.add_run(line)
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            run.font.size = Pt(14)
            run.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)
            paragraph.paragraph_format.space_before = Pt(13)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.25
            
            if section_match:
                section_num = section_match.group(1)
                if section_num in section_screenshot_map:
                    pending_screenshot = section_screenshot_map[section_num]

        else:
            paragraph.paragraph_format.first_line_indent = Pt(24)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.25

            parts = split_chinese_english(line)
            for text, is_chinese in parts:
                run = paragraph.add_run(text)
                run.font.size = Pt(12)
                run.font.color.rgb = RGBColor(0, 0, 0)
                if is_chinese:
                    run.font.name = '宋体'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                else:
                    run.font.name = 'Times New Roman'
                    run._element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')
                    run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    
    if pending_screenshot:
        section_title, screenshot_idx = pending_screenshot
        if screenshot_list and screenshot_idx < len(screenshot_list):
            insert_single_screenshot(doc, screenshot_list[screenshot_idx], section_title, screenshot_idx + 1)


def create_word_document(app_name, content, output_path):
    """
    创建Word文档 - 根据文件名判断文档类型
    """
    if "software_info" in output_path or "软著申请信息" in output_path:
        return create_software_info_document(app_name, content, output_path)
    elif "source_code" in output_path or "源代码" in output_path:
        return create_source_code_document(app_name, content, output_path)
    else:
        return create_manual_document(app_name, content, output_path)


def create_source_code_document(app_name, content, output_path):
    """
    创建源代码Word文档 - 完全基于ruanzhu19项目的原始实现
    """
    doc = Document()

    section = doc.sections[0]
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    setup_header_footer(doc, app_name, "source")

    lines = content.split('\n')

    for line in lines:
        paragraph = doc.add_paragraph()

        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0

        parts = split_chinese_english(line)
        for text, is_chinese in parts:
            run = paragraph.add_run(text)
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0, 0, 0)
            if is_chinese:
                run.font.name = '宋体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            else:
                run.font.name = 'Times New Roman'
                run._element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')
                run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Times New Roman')

    doc.save(output_path)
    return output_path


if __name__ == "__main__":
    
    screenshot_dir = "screenshots"

    if len(sys.argv) == 2:
        config_file = sys.argv[1]
        try:
            with open(config_file, 'r', encoding='utf-8') as f:
                config = json.load(f)
            
            app_name = config.get('appName', '')
            content = config.get('content', '')
            output_path = config.get('outputPath', '')

            if not app_name or not output_path:
                print(json.dumps({"error": "配置文件中缺少必要参数：appName 或 outputPath"}))
                sys.exit(1)

        except Exception as e:
            print(json.dumps({"error": f"读取配置文件失败: {str(e)}"}))
            sys.exit(1)

    elif len(sys.argv) == 4:
        app_name = sys.argv[1]
        content_param = sys.argv[2]
        output_path = sys.argv[3]
        screenshot_dir = "screenshots"
        
        if os.path.isfile(content_param):
            try:
                with open(content_param, 'r', encoding='utf-8') as f:
                    content = f.read()
            except UnicodeDecodeError:
                with open(content_param, 'r', encoding='gbk') as f:
                    content = f.read()
        else:
            content = content_param
    elif len(sys.argv) == 5:
        app_name = sys.argv[1]
        content_param = sys.argv[2]
        output_path = sys.argv[3]
        screenshot_dir = sys.argv[4]
        
        if os.path.isfile(content_param):
            try:
                with open(content_param, 'r', encoding='utf-8') as f:
                    content = f.read()
            except UnicodeDecodeError:
                with open(content_param, 'r', encoding='gbk') as f:
                    content = f.read()
        else:
            content = content_param
    elif len(sys.argv) == 6:
        app_name = sys.argv[1]
        content_param = sys.argv[2]
        output_path = sys.argv[3]
        screenshot_dir = sys.argv[4]
        backend_code_lines = sys.argv[5]
        
        if os.path.isfile(content_param):
            try:
                with open(content_param, 'r', encoding='utf-8') as f:
                    content = f.read()
            except UnicodeDecodeError:
                with open(content_param, 'r', encoding='gbk') as f:
                    content = f.read()
        else:
            content = content_param
        
        content += f"\n源代码总行数: {backend_code_lines}"
    else:
        print(json.dumps({"error": "参数错误，支持四种方式：1) 配置文件路径 2) 应用名称 内容 输出路径 3) 应用名称 内容 输出路径 截图目录 4) 应用名称 内容 输出路径 截图目录 源代码行数"}))
        sys.exit(1)

    try:
        if "software_info" in output_path:
            parsed_info = parse_software_info_content(content, app_name)
            try:
                sys.stderr.write("DEBUG: 软著申请信息文档生成完成\n")
                sys.stderr.write(f"DEBUG: 软件名称: {parsed_info['clean_name']}\n")
                sys.stderr.write(f"DEBUG: 源程序量: {parsed_info['code_lines']}\n")
                sys.stderr.write(f"DEBUG: 技术特点长度: {len(parsed_info['features'])} 字符\n")
                sys.stderr.flush()
            except Exception:
                pass

        result_path = create_word_document(app_name, content, output_path)
        print(json.dumps({"success": True, "outputPath": result_path}))
    except Exception as e:
        print(json.dumps({"error": str(e)}))
        sys.exit(1)
