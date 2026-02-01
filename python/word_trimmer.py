#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Word文档裁剪模块
"""

import sys
import os
import io
from docx import Document
import re
import subprocess
import platform
import time
import shutil
import math

if platform.system() == 'Windows':
    import msvcrt
else:
    import fcntl
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8')


class LibreOfficeLock:
    def __init__(self):
        self.lock_file = None
        self.lock_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), '.libreoffice.lock')
    
    def __enter__(self):
        self.lock_file = open(self.lock_path, 'w')
        max_wait = 60
        waited = 0
        
        while waited < max_wait:
            try:
                if platform.system() == 'Windows':
                    msvcrt.locking(self.lock_file.fileno(), msvcrt.LK_NBLCK, 1)
                else:
                    fcntl.flock(self.lock_file.fileno(), fcntl.LOCK_EX | fcntl.LOCK_NB)
                print(f"[LOCK] LibreOffice锁已获取")
                return self
            except (IOError, OSError):
                if waited == 0:
                    print(f"[LOCK] 等待其他LibreOffice任务完成...")
                time.sleep(2)
                waited += 2
        
        raise TimeoutError("无法获取LibreOffice锁，等待超时")
    
    def __exit__(self, exc_type, exc_val, exc_tb):
        if self.lock_file:
            try:
                if platform.system() == 'Windows':
                    msvcrt.locking(self.lock_file.fileno(), msvcrt.LK_UNLCK, 1)
                else:
                    fcntl.flock(self.lock_file.fileno(), fcntl.LOCK_UN)
                print(f"[LOCK] LibreOffice锁已释放")
                time.sleep(0.5)
            except:
                pass
            finally:
                self.lock_file.close()


def find_portable_libreoffice():
    system = platform.system()
    
    system_paths = []
    if system == 'Linux':
        system_paths = [
            '/usr/bin/libreoffice',
            '/usr/bin/soffice',
            '/usr/local/bin/libreoffice',
        ]
    elif system == 'Darwin':
        system_paths = [
            '/Applications/LibreOffice.app/Contents/MacOS/soffice',
        ]
    elif system == 'Windows':
        system_paths = [
            r'C:\Program Files\LibreOffice\program\soffice.exe',
            r'C:\Program Files (x86)\LibreOffice\program\soffice.exe',
        ]
    
    for path in system_paths:
        if os.path.exists(path):
            print(f"[INFO] 使用系统LibreOffice: {path}")
            return path
    
    script_dir = os.path.dirname(os.path.abspath(__file__))
    
    current = script_dir
    max_serve_root = None
    for _ in range(10):
        if os.path.basename(current) == 'max-serve':
            max_serve_root = current
            break
        parent = os.path.dirname(current)
        if parent == current:
            break
        current = parent
    
    portable_paths = []
    
    if max_serve_root:
        if system == 'Windows':
            portable_paths.extend([
                os.path.join(max_serve_root, 'App', 'libreoffice', 'program', 'soffice.exe'),
                os.path.join(max_serve_root, 'LibreOfficePortable', 'App', 'libreoffice', 'program', 'soffice.exe'),
                os.path.join(max_serve_root, 'libreoffice', 'program', 'soffice.exe'),
            ])
        elif system == 'Linux':
            portable_paths.extend([
                os.path.join(max_serve_root, 'libreoffice', 'program', 'soffice'),
                os.path.join(max_serve_root, 'LibreOffice', 'program', 'soffice'),
            ])
        elif system == 'Darwin':
            portable_paths.extend([
                os.path.join(max_serve_root, 'libreoffice', 'LibreOffice.app', 'Contents', 'MacOS', 'soffice'),
                os.path.join(max_serve_root, 'LibreOffice.app', 'Contents', 'MacOS', 'soffice'),
            ])
    
    for path in portable_paths:
        if os.path.exists(path):
            print(f"[INFO] 找到便携版LibreOffice: {path}")
            return path
    
    return None

def get_libreoffice_page_count(docx_file):
    with LibreOfficeLock():
        try:
            from PyPDF2 import PdfReader
        except ImportError:
            print("[错误] PyPDF2未安装，请运行: pip install PyPDF2")
            sys.exit(1)
        
        soffice_path = find_portable_libreoffice()
        if not soffice_path:
            print("[错误] 未找到LibreOffice便携版")
            print("[提示] 请下载LibreOffice便携版到 max-serve/App/libreoffice/")
            print("[下载] https://www.libreoffice.org/download/portable-versions/")
            sys.exit(1)
        
        print(f"[LibreOffice] 找到: {soffice_path}")
        print(f"[LibreOffice] 正在转换Word为PDF...")
        
        base_name = os.path.splitext(os.path.basename(docx_file))[0]
        output_dir = os.path.dirname(docx_file) or '.'
        pdf_file = os.path.join(output_dir, base_name + '.pdf')
        
        if os.path.exists(pdf_file):
            try:
                os.remove(pdf_file)
            except:
                pass
        
        cmd = [
            soffice_path,
            '--headless',
            '--convert-to', 'pdf',
            '--outdir', output_dir,
            docx_file
        ]
        
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=120)
        
        if result.returncode != 0 or not os.path.exists(pdf_file):
            print(f"[错误] PDF转换失败: {result.stderr}")
            raise RuntimeError("LibreOffice转换失败")
        
        try:
            pdf_reader = PdfReader(pdf_file)
            page_count = len(pdf_reader.pages)
            print(f"[LibreOffice] ✅ 精确页数: {page_count}页")
        except Exception as e:
            print(f"[错误] PDF读取失败: {e}")
            raise
        
        try:
            if os.path.exists(pdf_file):
                os.remove(pdf_file)
        except:
            pass
        
        return page_count

def get_word_page_count_pandoc(docx_file):
    try:
        import tempfile
        import shutil
        
        pandoc_check = subprocess.run(['pandoc', '--version'], 
                                     capture_output=True, text=True, timeout=5)
        if pandoc_check.returncode != 0:
            print(f"[PANDOC] Pandoc未安装，降级到LibreOffice")
            return None
        
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_pdf = os.path.join(temp_dir, 'output.pdf')
            abs_docx = os.path.abspath(docx_file)
            
            print(f"[PANDOC] 使用Pandoc + pdflatex转换...")
            
            cmd = [
                'pandoc',
                abs_docx,
                '-o', temp_pdf,
                '--pdf-engine=wkhtmltopdf'
            ]
            
            result = subprocess.run(cmd, capture_output=True, text=True, timeout=60)
            
            if result.returncode != 0:
                print(f"[PANDOC] 转换失败: {result.stderr[:200]}")
                return None
            
            if not os.path.exists(temp_pdf):
                print(f"[PANDOC] PDF文件未生成")
                return None
            
            try:
                from PyPDF2 import PdfReader
                reader = PdfReader(temp_pdf)
                page_count = len(reader.pages)
                print(f"[PANDOC] ✅ 精确页数: {page_count}页")
                return page_count
            except ImportError:
                pdfinfo_result = subprocess.run(['pdfinfo', temp_pdf],
                                               capture_output=True, text=True, timeout=5)
                if pdfinfo_result.returncode == 0:
                    for line in pdfinfo_result.stdout.split('\n'):
                        if line.startswith('Pages:'):
                            page_count = int(line.split(':')[1].strip())
                            print(f"[PANDOC] ✅ 精确页数: {page_count}页")
                            return page_count
                return None
                
    except subprocess.TimeoutExpired:
        print(f"[PANDOC] 转换超时")
        return None
    except Exception as e:
        print(f"[PANDOC] 错误: {str(e)}")
        return None

def get_word_page_count(docx_file):
    return get_word_page_count_libreoffice(docx_file)

def get_word_page_count_libreoffice(docx_file):
    return get_libreoffice_page_count(docx_file)

def remove_empty_paragraphs(docx_file):
    print(f"[REMOVE-EMPTY] 开始删除空行...")
    doc = Document(docx_file)
    
    indices_to_remove = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        is_empty = False
        
        if not text or len(text) == 0:
            is_empty = True
        elif all(c in '\r\n\t \u3000\xa0' for c in text):
            is_empty = True
        elif len(para.runs) == 0:
            is_empty = True
        elif all(not run.text.strip() for run in para.runs):
            is_empty = True
        
        if is_empty:
            indices_to_remove.append(i)
    
    removed_count = 0
    for idx in sorted(indices_to_remove, reverse=True):
        if idx < len(doc.paragraphs):
            para = doc.paragraphs[idx]
            p = para._element
            if p.getparent() is not None:
                p.getparent().remove(p)
                removed_count += 1
    
    doc.save(docx_file)
    print(f"[REMOVE-EMPTY] 已删除{removed_count}个空行")
    return removed_count

def remove_long_paragraphs(docx_file, max_chars=70):
    print(f"[REMOVE-LONG] 开始检查超长段落（>{max_chars}字符）...")
    
    total_removed = 0
    iteration = 0
    max_iterations = 10
    
    while iteration < max_iterations:
        iteration += 1
        doc = Document(docx_file)
        
        found_long = False
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if text and len(text) > max_chars:
                p = para._element
                parent = p.getparent()
                if parent is not None:
                    parent.remove(p)
                    doc.save(docx_file)
                    total_removed += 1
                    found_long = True
                    print(f"[REMOVE-LONG] 第{iteration}轮: 删除第{i}个段落({len(text)}字符)")
                    break
        
        if not found_long:
            print(f"[REMOVE-LONG] 迭代{iteration}轮，共删除{total_removed}个超长段落")
            break
    
    return total_removed

def get_paragraph_count(docx_file):
    doc = Document(docx_file)
    non_empty_count = len([p for p in doc.paragraphs if p.text.strip()])
    print(f"[PARAGRAPH-COUNT] 当前文档段落数: {non_empty_count}个")
    return non_empty_count

def trim_word_by_line_count(docx_file, lines_per_page=45, max_pages=60, target_lines=None):
    print(f"[TRIM-BY-PARAGRAPHS] ===== 开始裁剪 =====")
    
    if target_lines is None:
        target_lines = int(lines_per_page * max_pages)
    
    print(f"[TRIM-BY-PARAGRAPHS] 目标段落数: {target_lines}个")
    
    remove_long_paragraphs(docx_file, max_chars=70)
    remove_empty_paragraphs(docx_file)
    current_paragraphs = get_paragraph_count(docx_file)
    
    if current_paragraphs <= target_lines:
        print(f"[TRIM-BY-PARAGRAPHS] 当前{current_paragraphs}个段落 ≤ 目标{target_lines}个，无需裁剪")
        final_pages = get_word_page_count(docx_file)
        return final_pages
    
    paragraphs_to_remove = current_paragraphs - target_lines
    print(f"[TRIM-BY-PARAGRAPHS] 当前{current_paragraphs}个段落 > 目标{target_lines}个")
    print(f"[TRIM-BY-PARAGRAPHS] 需要删除{paragraphs_to_remove}个段落")
    
    doc = Document(docx_file)
    
    startup_indices = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if "if __name__ == '__main__':" in text or "app.run(debug=True)" in text:
            startup_indices.append(i)
            for offset in range(-5, 6):
                if 0 <= i + offset < len(doc.paragraphs):
                    startup_indices.append(i + offset)
    
    startup_indices = set(startup_indices)
    
    paragraphs_info = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        
        if i in startup_indices:
            priority = 0
        elif any(k in text for k in ['<html', '<div', '<script', 'function ', 'var ', 'const ']):
            priority = 0
        elif any(k in text for k in ['=== app.py ===', '=== api_routes.py ===', '@app.route']):
            priority = 0
        elif any(k in text for k in ['=== middleware.py ===', '@auth_required']):
            priority = 3
        elif any(k in text for k in ['=== config.py ===', '=== models.py ===']):
            priority = 5
        else:
            priority = 3
        
        paragraphs_info.append((i, para, priority, text))
    
    paragraphs_info.sort(key=lambda x: x[2], reverse=True)
    
    indices_to_delete = []
    deleted_count = 0
    
    for i, para, priority, text in paragraphs_info:
        if priority == 0:
            continue
        
        if deleted_count < paragraphs_to_remove:
            indices_to_delete.append(i)
            deleted_count += 1
        else:
            break
    
    indices_to_delete.sort(reverse=True)
    deleted_paras = 0
    
    for idx in indices_to_delete:
        if idx < len(doc.paragraphs):
            para = doc.paragraphs[idx]
            p = para._element
            if p.getparent() is not None:
                p.getparent().remove(p)
                deleted_paras += 1
    
    doc.save(docx_file)
    
    final_paragraphs = get_paragraph_count(docx_file)
    final_pages = get_word_page_count(docx_file)
    print(f"[TRIM-BY-PARAGRAPHS] 完成: {final_paragraphs}段落, {final_pages}页")
    
    return final_pages

def get_last_page_line_count(docx_file, target_page=60):
    pdf_file = None
    try:
        from PyPDF2 import PdfReader
        pdf_file = docx_file.replace('.docx', '.pdf')
        if os.path.exists(pdf_file):
            os.remove(pdf_file)
        
        with LibreOfficeLock():
            soffice_path = find_portable_libreoffice()
            if not soffice_path:
                print(f"[LINE-COUNT] 找不到LibreOffice")
                return 0
            
            print(f"[LINE-COUNT] 正在转换PDF以检查第{target_page}页...")
            cmd = [
                soffice_path,
                '--headless',
                '--convert-to', 'pdf',
                '--outdir', os.path.dirname(docx_file),
                docx_file
            ]
            
            result = subprocess.run(cmd, capture_output=True, text=True, timeout=60)
            
            if result.returncode != 0 or not os.path.exists(pdf_file):
                print(f"[LINE-COUNT] PDF转换失败")
                return 0
        
        reader = PdfReader(pdf_file)
        total_pages = len(reader.pages)
        
        if total_pages < target_page:
            print(f"[LINE-COUNT] PDF只有{total_pages}页，小于目标页{target_page}")
            os.remove(pdf_file)
            return 0
        
        page = reader.pages[target_page - 1]
        text = page.extract_text()
        
        lines = [line for line in text.split('\n') if line.strip()]
        line_count = len(lines)
        char_count = len(text)
        
        print(f"[LINE-COUNT] 第{target_page}页: {line_count}行, {char_count}字符")
        
        os.remove(pdf_file)
        
        return line_count
        
    except Exception as e:
        print(f"[LINE-COUNT] 错误: {e}")
        if pdf_file and os.path.exists(pdf_file):
            os.remove(pdf_file)
        return 0

def trim_word_to_pages_old(docx_file, max_pages, target_fill_ratio=0.9):
    doc = Document(docx_file)
    
    print(f"[STEP-1] 删除空行...")
    empty_count = 0
    for para in doc.paragraphs[:]:
        if not para.text.strip():
            p = para._element
            if p.getparent() is not None:
                p.getparent().remove(p)
                empty_count += 1
    
    if empty_count > 0:
        doc.save(docx_file)
        print(f"[STEP-1] 已删除{empty_count}个空行")
    else:
        print(f"[STEP-1] 无空行需要删除")
    
    doc = Document(docx_file)
    
    startup_code_indices = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if "if __name__ == '__main__':" in text or "app.run(debug=True)" in text:
            startup_code_indices.append(i)
            for offset in range(-5, 6):
                if 0 <= i + offset < len(doc.paragraphs):
                    startup_code_indices.append(i + offset)
    
    paragraphs_info = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        
        if i in startup_code_indices:
            priority = 0
        elif any(k in text for k in [
                '<html', '<div', '<head', '<body', '<script', '<style', 
                '<label', '<button', '<select', '<input', '<form', '<table',
                '<span', '<a ', '<img', '<ul', '<li', '<nav', '<header',
                '</html>', '</div>', '</head>', '</body>', '</script>', '</style>',
                '</label>', '</button>', '</select>', '</form>', '</table>',
                '</span>', '</a>', '</ul>', '</li>', '</nav>', '</header>',
                'class="', 'id="', '<!DOCTYPE', '<meta', '<link', '<section',
                'function ', 'var ', 'const ', 'let ', '$(', 'getElementById',
                'document.', 'addEventListener', 'querySelector', 'src="https://',
                'bootstrap', 'jquery', 'vue', 'react'
            ]):
            priority = 0
        elif any(k in text for k in ['=== services.py ===', '=== api_routes.py ===', '=== app.py ===',
                                      '@app.route', 'def create_app', 'Flask(__name__)']):
            priority = 0
        elif ('def get_' in text or 'def list_' in text or 'def update_' in text) and i > len(doc.paragraphs) * 0.7:
            priority = 2
        elif any(k in text for k in ['=== middleware.py ===', '=== validators.py ===', '=== auth.py ===']):
            priority = 3
        elif any(k in text for k in ['@auth_required', 'def validate_', 'class Validator', 'JWT']):
            priority = 3
        elif any(k in text for k in ['=== config.py ===', '=== models.py ===', '=== utils.py ===']):
            priority = 5
        elif any(k in text for k in ['class Config', 'DATABASE_URI', 'class Meta', 'db.Model']):
            priority = 5
        elif 'def delete_' in text or 'def format_' in text or 'def parse_' in text:
            priority = 4
        else:
            priority = 3
        
        paragraphs_info.append((i, para, priority, text))
    
    paragraphs_info.sort(key=lambda x: x[2], reverse=True)
    
    deleted = 0
    protected_count = len([x for x in paragraphs_info if x[2] == 0])
    deletable_paras = [(i, para, priority, text) for i, para, priority, text in paragraphs_info if priority > 0]
    
    print(f"[PHASE-1] 快速裁剪到接近{max_pages}页...")
    for i, para, priority, text in deletable_paras:
        if priority >= 4:
            p = para._element
            p.getparent().remove(p)
            deleted += 1
            
            if deleted % 20 == 0:
                temp_file = docx_file + '.tmp'
                doc.save(temp_file)
                current_pages = get_word_page_count(temp_file)
                print(f"[PHASE-1] 已删除Part1后端代码{deleted}段，当前{current_pages}页")
                
                if current_pages <= max_pages + 1:
                    os.remove(temp_file)
                    break
                os.remove(temp_file)
    
    doc.save(docx_file)
    current_pages = get_word_page_count(docx_file)
    print(f"[PHASE-2] 精确调整，当前{current_pages}页，目标{max_pages}页")
    
    if current_pages == max_pages:
        print(f"[OK] 页数正好符合要求：{max_pages}页")
        print(f"[OK] 已保留启动代码: if __name__ == '__main__': app.run(debug=True)")
        return max_pages
    
    if current_pages > max_pages:
        print(f"[PHASE-2] 页数过多（{current_pages}>{max_pages}），继续精确裁剪（删除Part2：middleware/auth）...")
        remaining_paras = [(i, para, priority, text) for i, para, priority, text in deletable_paras 
                          if priority == 3 and para._element.getparent() is not None]
        
        for idx, (i, para, priority, text) in enumerate(remaining_paras):
            p = para._element
            if p.getparent() is not None:
                p.getparent().remove(p)
                deleted += 1
                
                if current_pages > max_pages + 1:
                    check_interval = 30
                elif current_pages > max_pages:
                    check_interval = 1
                else:
                    print(f"[INFO] 当前{current_pages}页，已达到目标，停止删除")
                    break
                
                if (idx + 1) % check_interval == 0:
                    doc.save(docx_file)
                    retry_count = 0
                    max_retries = 3
                    while retry_count < max_retries:
                        try:
                            current_pages = get_word_page_count(docx_file)
                            break
                        except TimeoutError as e:
                            retry_count += 1
                            if retry_count < max_retries:
                                print(f"[WARN] 页数计算超时，重试 {retry_count}/{max_retries}...")
                                time.sleep(3)
                            else:
                                print(f"[ERROR] 页数计算失败，停止裁剪")
                                raise
                    print(f"[PHASE-2] 已删除Part2后端代码{deleted}段，当前{current_pages}页")
                    
                    if current_pages == max_pages:
                        print(f"[SUCCESS] 精确裁剪成功！正好{max_pages}页")
                        print(f"[OK] 已保留启动代码: if __name__ == '__main__': app.run(debug=True)")
                        return max_pages
                    elif current_pages < max_pages:
                        print(f"[WARN] 删除过多（{current_pages}<{max_pages}），停止删除...")
                        break
                    elif current_pages == max_pages - 1:
                        print(f"[INFO] 已接近目标页数（{current_pages}页），停止删除...")
                        break
    
    elif current_pages < max_pages:
        print(f"[PHASE-3] 页数不足（{current_pages}<{max_pages}），添加填充代码...")
        fill_needed = max_pages - current_pages
        
        fill_code_lines = []
        for i in range(fill_needed * 3):
            func_code = f"""
def utility_function_{i+1}(data):
    \"\"\"工具函数：数据处理和验证\"\"\"
    if not data:
        return None
    result = {{}}
    for key, value in data.items():
        if isinstance(value, str):
            result[key] = value.strip()
        elif isinstance(value, (int, float)):
            result[key] = value
    return result
"""
            fill_code_lines.append(func_code)
        
        fill_text = "\n".join(fill_code_lines)
        
        insert_index = len(doc.paragraphs) - 10
        if insert_index > 0:
            doc.paragraphs[insert_index].insert_paragraph_before(fill_text)
            doc.save(docx_file)
            current_pages = get_word_page_count(docx_file)
            print(f"[PHASE-3] 填充{fill_needed * 3}个工具函数后页数: {current_pages}页")
    
    doc.save(docx_file)
    final_pages = get_word_page_count(docx_file)
    print(f"[FINAL] 最终页数: {final_pages}页（目标: {max_pages}页）")
    print(f"[OK] 已保留启动代码: if __name__ == '__main__': app.run(debug=True)")
    return final_pages

def ensure_last_page_filled(docx_file, max_pages, target_lines=50):
    """确保最后一页填充到指定行数以上（防止最后一页空白）
    
    Args:
        docx_file: Word文件路径
        max_pages: 目标页数
        target_lines: 最后一页的目标行数（默认50行）
    """
    print(f"[FILL-START] ========== 开始填充函数 ==========")
    print(f"[FILL-START] 文件: {docx_file}")
    print(f"[FILL-START] 目标页数: {max_pages}, 目标行数: {target_lines}")
    
    doc = Document(docx_file)
    print(f"[FILL-START] Document对象已创建")
    
    current_pages = get_word_page_count(docx_file)
    print(f"[FILL-CHECK] 当前页数: {current_pages}, 目标: {max_pages}")
    print(f"[FILL-CHECK] 判断条件 current_pages <= max_pages: {current_pages} <= {max_pages} = {current_pages <= max_pages}")
    
    if current_pages <= max_pages:
        if current_pages == max_pages:
            print(f"[FILL-PHASE] 页数正好{max_pages}页，但需要填充最后一页防止空白...")
        else:
            print(f"[FILL-PHASE] 页数{current_pages}页，尝试填充到{max_pages}页...")
        
        fill_functions = [
            "def is_valid(x):\n    return x is not None",
            "def get_value(d, k):\n    return d.get(k) if isinstance(d, dict) else None",
            "def safe_int(x):\n    try: return int(x)\n    except: return 0",
            "def safe_float(x):\n    try: return float(x)\n    except: return 0.0",
            "def is_empty(x):\n    return x is None or x == ''",
            "def merge(a, b):\n    return {**a, **b} if isinstance(a, dict) else a",
            "def filter_list(lst):\n    return [x for x in lst if x is not None]",
            "def count_items(data):\n    return len(data) if hasattr(data, '__len__') else 0",
            "def format_str(s):\n    return str(s).strip() if s else ''",
            "def parse_bool(x):\n    return x in (True, 'true', '1', 'yes')",
            "def get_type(x):\n    return type(x).__name__",
            "def is_numeric(x):\n    return isinstance(x, (int, float))",
            "def sum_values(lst):\n    return sum(x for x in lst if isinstance(x, (int, float)))",
            "def max_value(lst):\n    return max((x for x in lst if isinstance(x, (int, float))), default=0)",
            "def min_value(lst):\n    return min((x for x in lst if isinstance(x, (int, float))), default=0)",
        ]
        
        def find_startup_code_index(doc):
            """使用正则表达式查找启动代码位置，容忍空格变化和单行格式"""
            pattern = re.compile(r"if\s+__name__\s*==\s*['\"]__main__['\"]", re.IGNORECASE)
            
            for i, para in enumerate(doc.paragraphs):
                if pattern.search(para.text):
                    print(f"[DEBUG] 找到启动代码在第{i}段: {para.text[:50]}...")
                    return i
            
            print(f"[WARN] 未找到启动代码，将在文档末尾前插入")
            return -1
        
        startup_index = find_startup_code_index(doc)
        
        if startup_index == -1:
            startup_index = max(len(doc.paragraphs) - 5, 0)
            print(f"[INFO] 未找到启动代码，在第{startup_index}段前插入填充代码")
        
        func_index = 0
        print(f"[FILL-LOOP] 准备开始循环，共{len(fill_functions)}个函数可用")
        print(f"[FILL-LOOP] 循环条件: func_index({func_index}) < len(fill_functions)({len(fill_functions)}) = {func_index < len(fill_functions)}")
        
        while func_index < len(fill_functions):
            print(f"[FILL-LOOP] ===== 循环第{func_index + 1}次开始 =====")
            backup_file = docx_file + '.backup'
            doc.save(backup_file)
            
            insert_position = min(startup_index, len(doc.paragraphs) - 1)
            if insert_position < 0:
                insert_position = 0
            
            print(f"[DEBUG] 在第{insert_position}段前插入函数{func_index + 1}")
            
            
            current_func = fill_functions[func_index]
            func_lines = current_func.split('\n')
            
            print(f"[DEBUG] 函数{func_index + 1}包含{len(func_lines)}行代码")
            
            if insert_position < len(doc.paragraphs):
                for line_idx in reversed(range(len(func_lines))):
                    line_text = func_lines[line_idx].strip()
                    if line_text:
                        new_paragraph = doc.paragraphs[insert_position].insert_paragraph_before(line_text)
                        
                        from docx.shared import Pt
                        
                        paragraph_format = new_paragraph.paragraph_format
                        paragraph_format.line_spacing = 1.0
                        paragraph_format.space_before = Pt(0)
                        paragraph_format.space_after = Pt(0)
                        
                        if new_paragraph.runs:
                            run = new_paragraph.runs[0]
                            font = run.font
                            font.name = 'Times New Roman'
                            font.size = Pt(12)
                        
                        print(f"[DEBUG] 插入行: {line_text[:30]}...")
            else:
                for line_text in func_lines:
                    line_text = line_text.strip()
                    if line_text:
                        new_paragraph = doc.add_paragraph(line_text)
                        
                        from docx.shared import Pt
                        
                        paragraph_format = new_paragraph.paragraph_format
                        paragraph_format.line_spacing = 1.0
                        paragraph_format.space_before = Pt(0)
                        paragraph_format.space_after = Pt(0)
                        
                        if new_paragraph.runs:
                            run = new_paragraph.runs[0]
                            font = run.font
                            font.name = 'Times New Roman'
                            font.size = Pt(12)
                        
                        print(f"[DEBUG] 追加行: {line_text[:30]}...")
            
            doc.save(docx_file)
            
            doc = Document(docx_file)
            startup_index = find_startup_code_index(doc)
            if startup_index == -1:
                startup_index = max(len(doc.paragraphs) - 5, 0)
            
            new_pages = get_word_page_count(docx_file)
            
            print(f"[FILL-PHASE] 添加函数{func_index + 1}，当前{new_pages}页，新startup_index={startup_index}")
            
            if new_pages > max_pages:
                print(f"[WARN] 页数超过目标（{new_pages}>{max_pages}），回退到上一个版本")
                shutil.copy2(backup_file, docx_file)
                os.remove(backup_file)
                doc = Document(docx_file)
                print(f"[INFO] 已回退，最终页数: {max_pages}页")
                break
            
            if new_pages == max_pages:
                line_count = get_last_page_line_count(docx_file, max_pages)
                if line_count >= target_lines:
                    print(f"[OK] 第{max_pages}页已有{line_count}行（目标{target_lines}行），停止填充")
                    os.remove(backup_file)
                    break
                else:
                    print(f"[INFO] 第{max_pages}页只有{line_count}行，继续填充...")
                    current_pages = new_pages
                    os.remove(backup_file)
                    func_index += 1
            else:
                print(f"[INFO] 继续填充到{max_pages}页...")
                current_pages = new_pages
                os.remove(backup_file)
                func_index += 1
        
        doc.save(docx_file)
        print(f"[SAVE] 最终版本已保存到: {docx_file}")
        
        final_pages = get_word_page_count(docx_file)
        final_line_count = get_last_page_line_count(docx_file, max_pages)
        print(f"[FILL-PHASE] 最终页数: {final_pages}页（目标: {max_pages}页）")
        print(f"[FILL-PHASE] 第{max_pages}页行数: {final_line_count}行（目标: {target_lines}行）")
        
        if final_line_count >= target_lines:
            print(f"[SUCCESS] 第{max_pages}页已填充到{final_line_count}行，达到目标{target_lines}行")
        else:
            print(f"[WARN] 第{max_pages}页只有{final_line_count}行，未达到目标{target_lines}行（可能是函数不够）")
    
    else:
        print(f"[WARN] 页数超过目标（{current_pages}>{max_pages}），无法添加填充")
    
    return

def ensure_last_page_filled_old(docx_file, max_pages):
    """旧版本：确保最后一页填充到90%以上"""
    doc = Document(docx_file)
    current_pages = get_word_page_count(docx_file)
    
    if current_pages != max_pages:
        print(f"[WARN] 页数不是目标值（{current_pages}≠{max_pages}），跳过填充检查")
        return
    
    print(f"[FILL-CHECK] 检查第{max_pages}页填充程度...")
    
    fill_lines = []
    helper_functions = [
        """
def validate_request_data(data, required_fields):
    \"\"\"验证请求数据完整性\"\"\"
    missing_fields = [field for field in required_fields if field not in data]
    if missing_fields:
        raise ValueError(f"缺少必填字段: {', '.join(missing_fields)}")
    return True
""",
        """
def format_response(status, message, data=None):
    \"\"\"格式化统一响应结构\"\"\"
    from datetime import datetime
    response = {
        'status': status,
        'message': message,
        'timestamp': datetime.now().isoformat()
    }
    if data is not None:
        response['data'] = data
    return response
""",
        """
def sanitize_input(input_str):
    \"\"\"清理和验证输入字符串\"\"\"
    if not isinstance(input_str, str):
        return str(input_str)
    cleaned = input_str.strip()
    cleaned = cleaned.replace('<', '&lt;').replace('>', '&gt;')
    return cleaned
"""
    ]
    fill_lines = helper_functions
    
    startup_index = -1
    for i, para in enumerate(doc.paragraphs):
        if "if __name__ == '__main__':" in para.text:
            startup_index = i
            break
    
    if startup_index > 0:
        insert_pos = startup_index - 1
        for func_code in fill_lines:
            doc.paragraphs[insert_pos].insert_paragraph_before(func_code.strip())
        
        doc.save(docx_file)
        new_pages = get_word_page_count(docx_file)
        
        if new_pages > max_pages:
            print(f"[FILL-ADJUST] 填充过多（{new_pages}页），减少1个辅助函数...")
            doc = Document(docx_file)
            delete_count = 0
            for i in range(startup_index - 1, max(0, startup_index - 10), -1):
                if i < len(doc.paragraphs):
                    text = doc.paragraphs[i].text.strip()
                    if 'def sanitize_input' in text or delete_count > 0:
                        p = doc.paragraphs[i]._element
                        if p.getparent() is not None:
                            p.getparent().remove(p)
                            delete_count += 1
                        if delete_count >= 8:
                            break
            doc.save(docx_file)
            final_pages = get_word_page_count(docx_file)
            print(f"[FILL-RESULT] 删除1个函数后页数: {final_pages}页")
        else:
            print(f"[FILL-RESULT] 添加{len(fill_lines)}个辅助函数，第{max_pages}页已优化填充")

def main():
    if len(sys.argv) < 3:
        print("用法: python word_trimmer.py <word文件> <最大页数>")
        sys.exit(1)
    
    word_file = sys.argv[1]
    max_pages = int(sys.argv[2])
    
    print(f"==== Word文档精确裁剪工具 ====")
    print(f"检查Word文档: {word_file}")
    current_pages = get_word_page_count(word_file)
    print(f"当前页数: {current_pages}页, 目标: {max_pages}页")
    
    if current_pages > max_pages:
        print(f"[WARN] 当前{current_pages}页 > 目标{max_pages}页，需要裁剪")
    else:
        print(f"[INFO] 当前{current_pages}页 <= 目标{max_pages}页，但仍需优化段落数")
    
    final_pages = trim_word_to_pages_old(word_file, max_pages)
    
    print(f"\n==== Word裁剪完成: {final_pages}页 ====")
    
    pdf_file = word_file.replace('.docx', '.pdf')
    if os.path.exists(pdf_file):
        os.remove(pdf_file)
    
    with LibreOfficeLock():
        soffice_path = find_portable_libreoffice()
        if soffice_path:
            print(f"[EXPORT] 正在转换为PDF...")
            cmd = [
                soffice_path,
                '--headless',
                '--convert-to', 'pdf',
                '--outdir', os.path.dirname(word_file) or '.',
                word_file
            ]
            result = subprocess.run(cmd, capture_output=True, text=True, timeout=120)
            
            if result.returncode == 0 and os.path.exists(pdf_file):
                from PyPDF2 import PdfReader
                pdf_reader = PdfReader(pdf_file)
                pdf_pages = len(pdf_reader.pages)
                print(f"[EXPORT] ✅ PDF格式: {pdf_file} ({pdf_pages}页)")
                
                if pdf_pages != max_pages:
                    print(f"[WARN] PDF页数({pdf_pages})与目标({max_pages})不符！")
            else:
                print(f"[ERROR] PDF转换失败: {result.stderr}")
        else:
            print(f"[ERROR] 未找到LibreOffice，无法生成PDF")
    
    print(f"\n==== 完成: {final_pages}页 ====")

if __name__ == '__main__':
    main()
